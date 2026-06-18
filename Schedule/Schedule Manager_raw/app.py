from flask import Flask, render_template, jsonify, request
import sqlite3
import os
import sys
import webbrowser
import threading
from flask import Flask, render_template, jsonify, request, send_file
from docxtpl import DocxTemplate
from io import BytesIO
from database import import_database
from contact_list import import_contact_list

# =========================
# BASE DIRECTORY (PYINSTALLER SAFE)
# =========================
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# =========================
# FLASK APP
# =========================
app = Flask(
    __name__,
    template_folder=os.path.join(BASE_DIR, "Templates")
)

# =========================
# PATHS
# =========================
DB_PATH = os.path.join(BASE_DIR, "database.db")


# =========================
# THEME COLOR REPLACEMENT
# =========================
def apply_theme_color(docx_path, old_hex, new_hex):
    """
    Replace every occurrence of old_hex (e.g. '002060') with new_hex
    inside a .docx file's internal XML (document, headers, footers).
    """
    import zipfile
    import shutil
    import tempfile

    old_hex = old_hex.upper()
    new_hex = new_hex.upper()

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:

                for item in zin.infolist():

                    content = zin.read(item.filename)

                    is_xml_target = (
                        item.filename.startswith("word/")
                        and item.filename.endswith(".xml")
                    )

                    if is_xml_target:
                        try:
                            text = content.decode("utf-8")
                            text = text.replace(old_hex, new_hex)
                            text = text.replace(old_hex.lower(), new_hex)
                            content = text.encode("utf-8")
                        except UnicodeDecodeError:
                            pass

                    zout.writestr(item, content)

        shutil.move(tmp_path, docx_path)

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


# =========================
# DB
# =========================
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


# =========================
# HOME (SHOW FULL DB)
# =========================
@app.route("/")
def home():
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM database')
    rows = cur.fetchall()
    conn.close()

    return render_template("index.html", data=rows)


# =========================
# SMART SEARCH (CN + EN + RANKING)
# =========================
@app.route("/search")
def search():
    q = request.args.get("q", "").strip()

    if not q:
        return jsonify([])

    conn = get_db()
    cur = conn.cursor()

    cur.execute("""
        SELECT *,
            CASE
                WHEN "机构名称（中文）" = ? THEN 1
                WHEN "机构名称（英文）" = ? THEN 1

                WHEN "机构名称（中文）" LIKE ? THEN 2
                WHEN "机构名称（英文）" LIKE ? THEN 2

                WHEN "机构名称（中文）" LIKE ? THEN 3
                WHEN "机构名称（英文）" LIKE ? THEN 3

                ELSE 4
            END as rank
        FROM database
        WHERE "机构名称（中文）" LIKE ?
           OR "机构名称（英文）" LIKE ?
        ORDER BY rank, "机构名称（中文）"
        LIMIT 30
    """, (
        q, q,              # exact
        f"{q}%", f"{q}%",  # prefix
        f"%{q}%", f"%{q}%",# contains
        f"%{q}%", f"%{q}%"
    ))

    results = cur.fetchall()
    conn.close()

    return jsonify([dict(r) for r in results])

@app.route("/search-contact")
def search_contact():

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    q = request.args.get("q", "").lower()
    q = " ".join(q.split())
    like = f"%{q}%"

    cur.execute("""
        SELECT *
        FROM Contact_list
        WHERE
            LOWER(REPLACE(REPLACE(IFNULL("English Full Name", ""), CHAR(160), ' '), '  ', ' ')) LIKE ?
            OR LOWER(REPLACE(REPLACE(IFNULL("Chinese Full Name", ""), CHAR(160), ' '), '  ', ' ')) LIKE ?
            OR LOWER(REPLACE(REPLACE(IFNULL("English Name", ""), CHAR(160), ' '), '  ', ' ')) LIKE ?
            OR LOWER(REPLACE(REPLACE(IFNULL("Work Title", ""), CHAR(160), ' '), '  ', ' ')) LIKE ?
            OR LOWER(REPLACE(REPLACE(IFNULL("Work Location", ""), CHAR(160), ' '), '  ', ' ')) LIKE ?
        LIMIT 20
    """, (like, like, like, like, like))

    rows = [dict(r) for r in cur.fetchall()]
    cur.execute("SELECT * FROM Contact_list LIMIT 1")
    print([desc[0] for desc in cur.description])
    conn.close()    

    return jsonify(rows)

@app.route("/debug-contact")
def debug_contact():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) FROM Contact_list")
    count = cur.fetchone()[0]

    cur.execute("SELECT * FROM Contact_list LIMIT 1")
    sample = cur.fetchone()

    return {
        "row_count": count,
        "sample_row": sample
    }
# =========================
# GENERATE WORD DOC
# =========================
import os
from flask import request, send_file
from docxtpl import DocxTemplate

@app.route("/generate-doc", methods=["POST"])
def generate_doc():
    try:
        # ================= PATH SETUP =================
        TEMPLATE_PATH = os.path.join(BASE_DIR, "template.docx")
        OUTPUT_PATH = os.path.join(BASE_DIR, "output.docx")
        TEMP_PATH = os.path.join(BASE_DIR, "temp.docx")

        if not os.path.exists(TEMPLATE_PATH):
            return "Template file not found", 500

        # ================= GET DATA =================
        data = request.get_json()

        if not data or "days" not in data:
            return "Invalid data format: missing 'days'", 400

        # ================= FORMAT DATE =================
        from datetime import datetime

        for day in data.get("days", []):

            raw_date = day.get("date")

            if raw_date:
                try:
                    dt = datetime.strptime(raw_date, "%Y-%m-%d")
                    day["date"] = f"{dt.month}月{dt.day}日"

                except:
                    pass

        # ================= ENRICH PARTICIPANTS =================
        for day in data.get("days", []):

            for meeting in day.get("meetings", []):

                for p in meeting.get("participants", []):

                    cn = (p.get("company_cn") or "").strip()
                    en = (p.get("company_en") or "").strip()

                    if not cn and not en:
                        p["formatted_companyName"] = ""

                    elif not en:
                        p["formatted_companyName"] = cn

                    elif not cn:
                        p["formatted_companyName"] = en

                    else:
                        p["formatted_companyName"] = f"{en}\n{cn}"

        # ================= BUILD ROWS =================

        # original flattened second table
        detail_table_rows = []

        # new participant table
        participantTable = []

        for day in data["days"]:

            for meeting in day["meetings"]:

                participants = meeting.get("participants", [])

                meeting["rows"] = []

                for i, p in enumerate(participants):

                    # ============================================
                    # FIRST TABLE
                    # ============================================
                    type_mode_language = ""

                    if i == 0:

                        parts = [
                            meeting.get("formatted_meetingtype", ""),
                            meeting.get("mode", ""),
                            meeting.get("language", "")
                        ]

                        type_mode_language = "\n".join(
                            [x for x in parts if x]
                        )

                    meeting["rows"].append({

                        "time":
                            meeting.get("time", "") if i == 0 else "",

                        "type_mode_language":
                            type_mode_language,

                        "formatted_companyName":
                            p.get("formatted_companyName", ""),

                        "type":
                            p.get("type", ""),

                        "contact":
                            p.get("contact", ""),

                        "address":
                            meeting.get("address", "") if i == 0 else "",

                        "bank":
                            "招证国际" if i == 0 else ""
                    })

                    # ============================================
                    # ORIGINAL DETAIL TABLE
                    # ============================================
                    detail_table_rows.append({

                        "time":
                            meeting.get("time", "") if i == 0 else "",

                        "meetingtype":
                            meeting.get("meetingtype", "") if i == 0 else "",

                        "language":
                            meeting.get("language", "") if i == 0 else "",

                        "formatted_companyName":
                            p.get("formatted_companyName", ""),

                        "type":
                            p.get("type", ""),

                        "aum":
                            p.get("aum", ""),

                        "contact":
                            p.get("contact", ""),

                        "intro":
                            p.get("intro", ""),

                        "case":
                            p.get("case", "")
                    })

                 # ================= NEW PARTICIPANT TABLE (DAY + PAGE CHUNKING) =================

                participantTable = []

                for day in data["days"]:

                    # 1. flatten all participants in this day
                    flat_rows = []

                    for meeting in day.get("meetings", []):
                        for p in meeting.get("participants", []):

                            flat_rows.append({
                                "formatted_companyName": p.get("formatted_companyName", ""),
                                "time": meeting.get("time", ""),
                                "meetingtype": meeting.get("meetingtype", ""),
                                "language": meeting.get("language", ""),
                                "type": p.get("type", ""),
                                "aum": p.get("aum", ""),
                                "contact": p.get("contact", ""),
                                "intro": p.get("intro", ""),
                                "case": p.get("case", "")
                            })

                    # 2. split into pages of max 3 participants
                    PAGE_SIZE = 3

                    for i in range(0, len(flat_rows), PAGE_SIZE):

                        participantTable.append({
                            "date": day.get("date", ""),
                            "location": day.get("location", ""),
                            "formatted_date": day.get("formatted_date", ""),
                            "rows": flat_rows[i:i + PAGE_SIZE]
                        })

        # ================= DOCX RENDER =================
        from docxtpl import DocxTemplate

        doc = DocxTemplate(TEMPLATE_PATH)

        context = {

            "projectName":
                data.get("projectName", ""),
                
            "typeEN":
                data.get("typeEN", ""),
            
            "typeCN":
                data.get("typeCN", ""),

            "startDate":
                data.get("startDate", ""),

            "formatted_startDate":
                data.get("formatted_startDate", ""),

            "endDate":
                data.get("endDate", ""),

            "formatted_endDate":
                data.get("formatted_endDate", ""),

            "managementTeam":
                data.get("managementTeam", []),

            "cmsTeam":
                data.get("cmsTeam", []),

            "location":
                data.get("location", ""),

            "formatted_location":
                data.get("formatted_location", ""),
            
            "period":
                data.get("period", ""),

            "days":
                data.get("days", []),

            # original detail table
            "detail_table_rows":
                detail_table_rows,

            # new participant table
            "participantTable":
                participantTable
        }

        # ================= DEBUG OUTPUT =================
        import json

        doc.render(context)

        # ================= MARGINS =================
        from docx.shared import Inches

        for section in doc.docx.sections:

            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        doc.save(TEMP_PATH)

        # ================= POST PROCESS =================
        from docx import Document

        document = Document(TEMP_PATH)

        from docx.oxml.ns import qn

        # ================= HELPER FUNCTIONS =================
        def is_row_empty(row):
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text and p.text.strip():
                        return False
            return True


        def clean_cell_empty_lines(cell):
            for p in list(cell.paragraphs):
                txt = p.text.replace("\xa0", "").strip()
                if txt == "":
                    p._element.getparent().remove(p._element)


        def is_participant_table(table):
            """Detect participant/detail tables so we SKIP merging"""
            text = " ".join(cell.text for row in table.rows for cell in row.cells)

            # your participant table indicators
            keywords = [
                "formatted_companyName",
                "aum",
                "intro",
                "case"
            ]

            return any(k in text for k in keywords)


        def is_merge_table(table):
            """Only allow merging for schedule/management/CMS tables"""
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells)
                if "线下地址 / 线上拨入" in row_text:
                    return True
            return False

        # ================= REMOVE EMPTY ROWS (ALL TABLES OK) =================
        for table in document.tables:

            rows = list(table.rows)

            for row in rows:
                if row._index == 0:
                    continue

                if is_row_empty(row):
                    tbl = row._element.getparent()
                    tbl.remove(row._element)

        # =====================================================
        # COLLECT ONLY SCHEDULE TABLES
        # =====================================================
        schedule_tables = []

        for table in document.tables:

            if is_participant_table(table):
                continue

            if not is_merge_table(table):
                continue

            schedule_tables.append(table)

        # =====================================================
        # EACH DAY USES ITS OWN TABLE
        # =====================================================
        for day_index, day in enumerate(data["days"]):

            if day_index >= len(schedule_tables):
                break

            table = schedule_tables[day_index]

            print(f"✅ Processing Day {day_index + 1} table")

            # =====================================================
            # FIND HEADER ROW
            # =====================================================
            header_row_index = None

            for i, row in enumerate(table.rows):

                row_text = " ".join(
                    cell.text.strip()
                    for cell in row.cells
                )

                if "负责投行" in row_text:
                    header_row_index = i
                    break

            if header_row_index is None:
                continue

            # =====================================================
            # RESET ROW INDEX FOR THIS DAY
            # =====================================================
            current_row = header_row_index + 1

            # =====================================================
            # ONLY PROCESS THIS DAY'S MEETINGS
            # =====================================================
            for meeting in day["meetings"]:

                participants = meeting.get("participants", [])

                participant_count = meeting.get("participantCount", 0)

                if not participant_count:
                    participant_count = max(len(participants), 1)

                start_row = current_row
                end_row = start_row + participant_count - 1

                print(
                    f"Meeting rows = {start_row} -> {end_row} "
                    f"(participants={participant_count})"
                )

                # =====================================================
                # SAFETY CHECK
                # =====================================================
                if end_row >= len(table.rows):

                    print(
                        f"⚠️ Skip merge: "
                        f"end_row={end_row}, "
                        f"table_rows={len(table.rows)}"
                    )

                    break

                # =====================================================
                # MERGE ONLY SAME MEETING ROWS
                # =====================================================
                if participant_count > 1:

                    for col in [0, 1, 6, 7]:

                        try:

                            top_cell = table.cell(start_row, col)
                            bottom_cell = table.cell(end_row, col)

                            merged_cell = top_cell.merge(bottom_cell)

                            clean_cell_empty_lines(merged_cell)

                            # preserve formatting
                            for p in merged_cell.paragraphs:

                                for run in p.runs:

                                    r = run._element
                                    rPr = r.get_or_add_rPr()
                                    rFonts = rPr.get_or_add_rFonts()

                                    rFonts.set(qn("w:ascii"), "Arial")
                                    rFonts.set(qn("w:hAnsi"), "Arial")
                                    rFonts.set(qn("w:eastAsia"), "KaiTi")

                            print(
                                f"✅ Merged col={col} "
                                f"rows={start_row}-{end_row}"
                            )

                        except Exception as e:

                            print(
                                f"❌ Merge failed "
                                f"col={col}: {e}"
                            )

                # =====================================================
                # MOVE TO NEXT MEETING
                # =====================================================
                current_row = end_row + 1

        # ================= REMOVE LAST PAGE BREAK =================
        from docx.enum.text import WD_BREAK

        def remove_last_page_break(doc):

            paragraphs = doc.paragraphs

            # traverse backwards
            for p in reversed(paragraphs):

                # inspect runs backwards
                for run in reversed(p.runs):

                    brs = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
                    )

                    if brs:

                        last_br = brs[-1]

                        # remove only PAGE breaks
                        br_type = last_br.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                        )

                        if br_type == "page":

                            run._element.remove(last_br)

                            return

        remove_last_page_break(document)

        # ================= SAVE FINAL =================
        document.save(OUTPUT_PATH)

        # ================= APPLY THEME COLOR =================
        theme_color = (data.get("themeColor") or "").strip().upper()

        # Only replace if it's a valid 6-digit hex color and different from default
        import re

        if re.fullmatch(r"[0-9A-F]{6}", theme_color) and theme_color != "002060":
            apply_theme_color(OUTPUT_PATH, "002060", theme_color)

        return send_file(
            OUTPUT_PATH,
            as_attachment=True,
            download_name="schedule.docx"
        )

    except Exception as e:

        import traceback
        traceback.print_exc()

        return str(e), 500
    
# =========================
# AUTO OPEN
# =========================
def open_browser():
    webbrowser.open("http://127.0.0.1:5000")

def start_import():
    import_database()
    import_contact_list()

if __name__ == "__main__":
    threading.Thread(target=start_import).start()
    threading.Timer(1.2, open_browser).start()
    app.run(
        host="127.0.0.1",
        port=5000,
        debug=False,
        use_reloader=False
    )